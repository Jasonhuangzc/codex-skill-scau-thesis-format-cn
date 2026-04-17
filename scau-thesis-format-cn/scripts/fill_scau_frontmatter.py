#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


FIELD_RUN_EXPECTATIONS = {
    1: 0,
    3: 0,
    10: 6,
    11: 5,
    12: 5,
    13: 5,
    14: 9,
    15: 9,
    39: 0,
    41: 0,
    42: 0,
    43: 0,
    44: 0,
    45: 0,
    47: 0,
}

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
BUNDLED_TEMPLATE_DOCX = SKILL_ROOT / "assets" / "template" / "scau-undergrad-thesis-template.docx"


def discover_workspace_root(start: Path) -> Path:
    for candidate in [start, *start.parents]:
        if (candidate / "thesis_metadata.json").exists():
            return candidate
        generic_meta = list(candidate.glob("*metadata*.json"))
        if generic_meta:
            return candidate
    return start


def discover_meta(workspace: Path) -> Path:
    meta_path = workspace / "thesis_metadata.json"
    if meta_path.exists():
        return meta_path
    generic_meta = sorted(workspace.rglob("*metadata*.json"))
    if generic_meta:
        return generic_meta[0]
    matches = list(workspace.rglob("thesis_metadata.json"))
    if not matches:
        raise FileNotFoundError("Could not find thesis_metadata.json.")
    return matches[0]


def discover_template(workspace: Path) -> Path:
    preferred = workspace / "论文撰写规范" / "附件6_格式模板_转存.docx"
    if preferred.exists():
        return preferred

    candidates = sorted(workspace.rglob("*格式模板*.docx"))
    if candidates:
        return candidates[0]

    english_named = sorted(workspace.rglob("*template*.docx"))
    if english_named:
        return english_named[0]

    if BUNDLED_TEMPLATE_DOCX.exists():
        return BUNDLED_TEMPLATE_DOCX
    raise FileNotFoundError(
        "Could not find a converted .docx thesis template. In the public repo, run scripts/import_official_2024_assets.py or pass --template explicitly."
    )


def discover_work_output_dir(workspace: Path) -> Path:
    for candidate in (
        workspace / "论文终稿",
        workspace / "work",
        workspace / "output",
        workspace / "outputs",
    ):
        if candidate.exists() and candidate.is_dir():
            return candidate
    return workspace / "_scau_thesis_output"


def default_output_path_for_workspace(workspace: Path) -> Path:
    output_dir = discover_work_output_dir(workspace)
    if output_dir.name == "论文终稿":
        return output_dir / "毕业论文终稿_工作版.docx"
    return output_dir / "scau_thesis_working.docx"


def load_metadata(meta_path: Path) -> dict:
    return json.loads(meta_path.read_text(encoding="utf-8"))


def coerce_text(value: object, default: str) -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def coerce_keywords(value: object, *, sep: str, default: str) -> str:
    if value is None:
        return default
    if isinstance(value, list):
        parts = [str(item).strip() for item in value if str(item).strip()]
        return sep.join(parts) if parts else default
    text = str(value).strip()
    return text if text else default


def set_run_font(
    run,
    east_asia_font: str,
    ascii_font: str,
    *,
    size_pt: float | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = ascii_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    run_pr = run._element.get_or_add_rPr()
    run_fonts = run_pr.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_pr.insert(0, run_fonts)
    run_fonts.set(qn("w:ascii"), ascii_font)
    run_fonts.set(qn("w:hAnsi"), ascii_font)
    run_fonts.set(qn("w:cs"), ascii_font)
    run_fonts.set(qn("w:eastAsia"), east_asia_font)


def fill_run(
    paragraph,
    run_index: int,
    text: str,
    east_asia_font: str = "宋体",
    ascii_font: str = "Times New Roman",
    *,
    size_pt: float | None = None,
    bold: bool | None = None,
) -> None:
    if run_index >= len(paragraph.runs):
        raise IndexError(f"Run index {run_index} out of range for paragraph: {paragraph.text!r}")
    run = paragraph.runs[run_index]
    run.text = text
    set_run_font(run, east_asia_font, ascii_font, size_pt=size_pt, bold=bold)


def clear_other_runs(paragraph, keep_indices: set[int]) -> None:
    for idx, run in enumerate(paragraph.runs):
        if idx not in keep_indices:
            run.text = ""


def reset_paragraph_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def set_label_body_runs(
    paragraph,
    *,
    label: str,
    body: str,
    label_east_asia: str,
    label_ascii: str,
    body_east_asia: str,
    body_ascii: str,
    label_size_pt: float | None = None,
    body_size_pt: float | None = None,
    label_bold: bool | None = None,
    body_bold: bool | None = None,
) -> None:
    reset_paragraph_runs(paragraph)
    if paragraph.runs:
        label_run = paragraph.runs[0]
    else:
        label_run = paragraph.add_run()
    label_run.text = label
    set_run_font(label_run, label_east_asia, label_ascii, size_pt=label_size_pt, bold=label_bold)

    body_run = paragraph.add_run(body)
    set_run_font(body_run, body_east_asia, body_ascii, size_pt=body_size_pt, bold=body_bold)


def split_submission_date(text: str) -> tuple[str, str, str]:
    parts = re.findall(r"\d+", text)
    if len(parts) < 3:
        raise ValueError(f"Invalid submission date format: {text}")
    return parts[0], parts[1], parts[2]


def validate_template(doc: Document) -> None:
    paragraphs = doc.paragraphs
    for index, required_run_index in FIELD_RUN_EXPECTATIONS.items():
        if index >= len(paragraphs):
            raise RuntimeError(f"Template does not contain paragraph index {index}.")
        if required_run_index >= len(paragraphs[index].runs):
            raise RuntimeError(
                f"Template structure mismatch at paragraph {index}: "
                f"expected run index {required_run_index}, got only {len(paragraphs[index].runs)} runs."
            )


def replace_cover(doc: Document, meta: dict, paper_type: str) -> None:
    paragraphs = doc.paragraphs
    fill_run(paragraphs[1], 0, paper_type, size_pt=36, bold=True)
    clear_other_runs(paragraphs[1], {0})

    fill_run(paragraphs[3], 0, meta["thesis_title_zh"], east_asia_font="黑体", size_pt=22, bold=True)
    clear_other_runs(paragraphs[3], {0})

    fill_run(paragraphs[10], 6, meta["college"], size_pt=15, bold=False)
    fill_run(paragraphs[11], 5, meta["major"], size_pt=15, bold=False)
    fill_run(paragraphs[12], 5, meta["student_name_zh"], size_pt=15, bold=False)
    fill_run(paragraphs[13], 5, meta["student_id"], east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=False)
    fill_run(paragraphs[14], 4, meta["advisor_name_zh"], size_pt=15, bold=False)
    fill_run(paragraphs[14], 9, meta["advisor_title"], size_pt=15, bold=False)

    year, month, day = split_submission_date(meta["submission_date"])
    fill_run(paragraphs[15], 2, year, east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=False)
    fill_run(paragraphs[15], 5, month, east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=False)
    fill_run(paragraphs[15], 9, day, east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=False)


def replace_abstract_frontmatter(doc: Document, meta: dict) -> None:
    paragraphs = doc.paragraphs

    zh_abstract = coerce_text(meta.get("abstract_zh"), "[中文摘要待写入]")
    zh_keywords = coerce_keywords(meta.get("keywords_zh"), sep="；", default="待补充；待补充；待补充")
    en_abstract = coerce_text(meta.get("abstract_en"), "[English abstract to be added]")
    en_keywords = coerce_keywords(meta.get("keywords_en"), sep="; ", default="To be added; To be added; To be added")

    fill_run(paragraphs[39], 0, zh_abstract, size_pt=12, bold=False)
    clear_other_runs(paragraphs[39], {0})
    clear_other_runs(paragraphs[40], set())
    set_label_body_runs(
        paragraphs[41],
        label="关键词：",
        body=zh_keywords,
        label_east_asia="黑体",
        label_ascii="Times New Roman",
        body_east_asia="宋体",
        body_ascii="Times New Roman",
        label_size_pt=12,
        body_size_pt=12,
        label_bold=False,
        body_bold=False,
    )

    fill_run(paragraphs[42], 0, meta["thesis_title_en"], east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=14, bold=True)
    clear_other_runs(paragraphs[42], {0})
    fill_run(paragraphs[43], 0, meta["english_name"], east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=12, bold=False)
    clear_other_runs(paragraphs[43], {0})
    affiliation = (
        f"（{meta['college_en']}, {meta['university_en']}, "
        f"{meta['city_en']} {meta['postal_code']}, China）"
    )
    fill_run(paragraphs[44], 0, affiliation, east_asia_font="Times New Roman", ascii_font="Times New Roman", size_pt=12, bold=False)
    clear_other_runs(paragraphs[44], {0})
    set_label_body_runs(
        paragraphs[45],
        label="Abstract:",
        body=f" {en_abstract}" if en_abstract and not en_abstract.startswith(" ") else en_abstract,
        label_east_asia="Times New Roman",
        label_ascii="Times New Roman",
        body_east_asia="Times New Roman",
        body_ascii="Times New Roman",
        label_size_pt=12,
        body_size_pt=12,
        label_bold=True,
        body_bold=False,
    )
    clear_other_runs(paragraphs[46], set())
    set_label_body_runs(
        paragraphs[47],
        label="Key words:",
        body=f" {en_keywords}" if en_keywords and not en_keywords.startswith(" ") else en_keywords,
        label_east_asia="Times New Roman",
        label_ascii="Times New Roman",
        body_east_asia="Times New Roman",
        body_ascii="Times New Roman",
        label_size_pt=12,
        body_size_pt=12,
        label_bold=True,
        body_bold=False,
    )


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def paragraph_has_page_break(paragraph) -> bool:
    xml = paragraph._p.xml
    return 'w:type="page"' in xml or "<w:lastRenderedPageBreak" in xml


def find_paragraph_index_startswith(paragraphs, prefix: str) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    return None


def ensure_frontmatter_page_breaks(doc: Document) -> dict:
    paragraphs = doc.paragraphs
    abstract_index = find_paragraph_index_startswith(paragraphs, "Abstract:")
    if abstract_index is None:
        return {"english_abstract_page_break": "missing_abstract_label"}

    english_title_index = None
    previous_indices: list[int] = []
    current = abstract_index - 1
    while current >= 0 and len(previous_indices) < 3:
        if paragraphs[current].text.strip():
            previous_indices.append(current)
        current -= 1
    previous_indices.reverse()
    if previous_indices:
        english_title_index = previous_indices[0]

    if english_title_index is None or english_title_index >= len(paragraphs):
        return {"english_abstract_page_break": "missing_title_target"}
    paragraphs[english_title_index].paragraph_format.page_break_before = True
    return {
        "english_abstract_page_break": "applied",
        "paragraph_index": english_title_index,
    }


def normalize_cover_gap(doc: Document) -> dict:
    paragraphs = doc.paragraphs
    declaration_index = None
    for index, paragraph in enumerate(paragraphs):
        if "原创性声明" in paragraph.text:
            declaration_index = index
            break
    if declaration_index is None:
        return {"declaration_found": False, "removed_blank_paragraphs": 0, "kept_page_break_paragraphs": 0}

    cover_end_index = 15
    mid_paragraphs = list(paragraphs[cover_end_index + 1 : declaration_index])
    removed = 0
    kept_breaks = 0
    for paragraph in mid_paragraphs:
        text = paragraph.text.strip()
        has_page_break = paragraph_has_page_break(paragraph)
        if has_page_break:
            kept_breaks += 1
            continue
        if not text:
            remove_paragraph(paragraph)
            removed += 1
    return {
        "declaration_found": True,
        "declaration_index": declaration_index,
        "removed_blank_paragraphs": removed,
        "kept_page_break_paragraphs": kept_breaks,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill the South China Agricultural University thesis front matter from a metadata JSON file.")
    parser.add_argument("--workspace", help="Workspace root. Defaults to current directory or its parents.")
    parser.add_argument("--meta", help="Path to metadata JSON. Supports thesis_metadata.json and generic metadata.json names.")
    parser.add_argument("--template", help="Path to the converted template .docx")
    parser.add_argument("--output", help="Path to output .docx. Defaults to a detected working directory, then falls back to _scau_thesis_output/scau_thesis_working.docx")
    parser.add_argument("--paper-type", default="本科毕业论文", help="Either 本科毕业论文 or 本科毕业设计")
    args = parser.parse_args()

    start = Path(args.workspace).resolve() if args.workspace else Path.cwd().resolve()
    workspace = discover_workspace_root(start)
    meta_path = Path(args.meta).resolve() if args.meta else discover_meta(workspace)
    template_path = Path(args.template).resolve() if args.template else discover_template(workspace)
    output_path = (
        Path(args.output).resolve()
        if args.output
        else default_output_path_for_workspace(workspace)
    )

    meta = load_metadata(meta_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(template_path)
    validate_template(doc)
    replace_cover(doc, meta, args.paper_type)
    replace_abstract_frontmatter(doc, meta)
    cover_gap_report = normalize_cover_gap(doc)
    frontmatter_page_breaks = ensure_frontmatter_page_breaks(doc)
    doc.save(output_path)
    print(
        json.dumps(
            {
                "output": str(output_path),
                "frontmatter_checks": {
                    "cover_to_declaration": cover_gap_report,
                    "frontmatter_page_breaks": frontmatter_page_breaks,
                },
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
