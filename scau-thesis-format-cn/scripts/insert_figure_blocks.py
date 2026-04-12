#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import math
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm
from PIL import Image

from compose_panel_grid import compose_from_spec

from word_template_utils import (
    default_output_path,
    find_heading_donors,
    insert_paragraph_after,
    insert_paragraph_before,
)


def load_manifest(path: Path) -> list[dict]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        raise ValueError("Figure manifest must be a JSON array.")
    return data


def resolve_image_paths(entry: dict, base_dir: Path) -> list[Path]:
    layout = (entry.get("layout") or "").lower()
    if layout in {"grid", "panel-grid"} or "panels" in entry:
        output_path = None
        if entry.get("composite_output"):
            output_path = Path(entry["composite_output"])
            if not output_path.is_absolute():
                output_path = (base_dir / output_path).resolve()
        else:
            temp_dir = Path(tempfile.gettempdir()) / "thesis-word-template-cn"
            temp_dir.mkdir(parents=True, exist_ok=True)
            stem = re.sub(r"\s+", "_", entry.get("caption", "panel_grid"))[:80]
            output_path = temp_dir / f"{stem}.png"
        return [compose_from_spec(entry, base_dir, output_path=output_path)]

    if "image" in entry:
        raw_paths = [entry["image"]]
    else:
        raw_paths = entry.get("images", [])
    if not raw_paths:
        raise ValueError("Each figure entry must provide image or images.")

    resolved: list[Path] = []
    for raw in raw_paths:
        path = Path(raw)
        if not path.is_absolute():
            path = (base_dir / path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"Image not found: {path}")
        resolved.append(path)
    return resolved


def find_anchor(document: Document, pattern: str, occurrence: int = 1):
    regex = re.compile(pattern)
    hits = []
    for paragraph in document.paragraphs:
        if regex.search(paragraph.text):
            hits.append(paragraph)
    if len(hits) < occurrence:
        raise RuntimeError(f"Anchor regex {pattern!r} matched {len(hits)} paragraphs, need {occurrence}.")
    return hits[occurrence - 1]


def set_picture_paragraph_defaults(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def insert_page_break_after(block, donor):
    page_break_para = insert_paragraph_after(block, donor, "")
    run = page_break_para.runs[0] if page_break_para.runs else page_break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    page_break_para.paragraph_format.keep_with_next = True
    return page_break_para


def insert_page_break_before(block, donor):
    page_break_para = insert_paragraph_before(block, donor, "")
    run = page_break_para.runs[0] if page_break_para.runs else page_break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    page_break_para.paragraph_format.keep_with_next = True
    return page_break_para


def estimate_image_block_height_cm(image_paths: list[Path], width_cm: float) -> float:
    total_height = 0.0
    for image_path in image_paths:
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        if width_px <= 0:
            continue
        total_height += width_cm * (height_px / width_px)
    if len(image_paths) > 1:
        total_height += 0.35 * (len(image_paths) - 1)
    return total_height


def estimate_note_height_cm(note: str | None, width_cm: float) -> float:
    if not note:
        return 0.0
    chars_per_line = max(18, int(width_cm * 4.8))
    lines = max(1, math.ceil(len(note) / chars_per_line))
    return lines * 0.48


def should_insert_page_break(entry: dict, images: list[Path], width_cm: float) -> bool:
    if entry.get("page_break_before") is True:
        return True
    if entry.get("page_break_before") is False:
        return False

    image_height = estimate_image_block_height_cm(images, width_cm)
    caption_height = 0.8
    note_height = estimate_note_height_cm(entry.get("note"), width_cm)
    estimated_total = image_height + caption_height + note_height + 0.8

    if estimated_total >= 11.2:
        return True
    if note_height >= 2.2 and image_height >= 6.0:
        return True
    return False


def insert_entry_after(anchor, entry: dict, donors: dict[str, object], base_dir: Path):
    images = resolve_image_paths(entry, base_dir)
    width_cm = float(entry.get("width_cm", 12.0))
    if should_insert_page_break(entry, images, width_cm):
        anchor = insert_page_break_after(anchor, donors["body"])
    cursor = insert_paragraph_after(anchor, donors["body"], "")
    set_picture_paragraph_defaults(cursor)
    for image_path in images:
        run = cursor.runs[0] if cursor.runs else cursor.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        if image_path != images[-1]:
            cursor = insert_paragraph_after(cursor, donors["body"], "")
            set_picture_paragraph_defaults(cursor)

    cursor = insert_paragraph_after(cursor, donors["figure_caption"], entry["caption"])
    cursor.paragraph_format.keep_with_next = bool(entry.get("note"))
    cursor.paragraph_format.keep_together = True
    if entry.get("note"):
        cursor = insert_paragraph_after(cursor, donors["note"], entry["note"])
        cursor.paragraph_format.keep_together = True
    return insert_paragraph_after(cursor, donors["body"], "")


def insert_entry_before(anchor, entry: dict, donors: dict[str, object], base_dir: Path):
    images = resolve_image_paths(entry, base_dir)
    width_cm = float(entry.get("width_cm", 12.0))
    if should_insert_page_break(entry, images, width_cm):
        anchor = insert_page_break_before(anchor, donors["body"])

    spacer = insert_paragraph_before(anchor, donors["body"], "")
    current_anchor = spacer

    if entry.get("note"):
        current_anchor = insert_paragraph_before(current_anchor, donors["note"], entry["note"])
    current_anchor = insert_paragraph_before(current_anchor, donors["figure_caption"], entry["caption"])

    for image_path in reversed(images):
        image_para = insert_paragraph_before(current_anchor, donors["body"], "")
        set_picture_paragraph_defaults(image_para)
        run = image_para.runs[0] if image_para.runs else image_para.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        current_anchor = image_para

    return insert_paragraph_before(current_anchor, donors["body"], "")


def main() -> None:
    parser = argparse.ArgumentParser(description="Insert figure blocks into a thesis .docx file from a JSON manifest.")
    parser.add_argument("--docx", required=True, help="Path to the input .docx")
    parser.add_argument("--manifest", required=True, help="Path to the figure manifest JSON")
    parser.add_argument("--fallback-template-docx", help="Optional official template .docx used when the working copy lacks donor paragraphs.")
    parser.add_argument("--output", help="Output .docx path. Defaults to overwrite-safe sibling path.")
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    manifest_path = Path(args.manifest).resolve()
    output_path = Path(args.output).resolve() if args.output else default_output_path(docx_path, "_插图插入")

    document = Document(docx_path)
    entries = load_manifest(manifest_path)
    fallback_document = None
    if args.fallback_template_docx:
        fallback_document = Document(Path(args.fallback_template_docx).resolve())
    required_donors = ["body", "figure_caption"]
    if any(entry.get("note") for entry in entries):
        required_donors.append("note")
    try:
        donors = find_heading_donors(
            document,
            required=required_donors,
            fallback_document=fallback_document,
        )
    except RuntimeError as exc:
        raise RuntimeError(
            "Figure donor parsing failed. "
            "Current flow only requires body/figure_caption/note donors for figure-only insertion. "
            "If the working copy has already removed template sample figure captions, pass --fallback-template-docx "
            "with the official template .docx so donor styles can be loaded from there."
        ) from exc
    base_dir = manifest_path.parent

    for entry in entries:
        anchor = find_anchor(
            document,
            entry["anchor_regex"],
            occurrence=int(entry.get("occurrence", 1)),
        )
        position = entry.get("position", "after").lower()
        if position == "after":
            insert_entry_after(anchor, entry, donors, base_dir)
        elif position == "before":
            insert_entry_before(anchor, entry, donors, base_dir)
        else:
            raise ValueError(f"Unsupported figure insertion position: {position}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(
        json.dumps(
            {"output": str(output_path), "manifest": str(manifest_path), "entries": len(entries)},
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
