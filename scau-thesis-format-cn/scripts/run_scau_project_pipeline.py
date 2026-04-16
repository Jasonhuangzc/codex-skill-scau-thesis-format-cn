#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
BUNDLED_TEMPLATE_DOCX = SKILL_ROOT / "assets" / "template" / "scau-undergrad-thesis-template.docx"


class SkillStepError(RuntimeError):
    pass


def emit_text(text: str, *, stderr: bool = False) -> None:
    stream = sys.stderr if stderr else sys.stdout
    payload = text if text.endswith("\n") else text + "\n"
    encoding = getattr(stream, "encoding", None) or "utf-8"
    stream.buffer.write(payload.encode(encoding, errors="backslashreplace"))
    stream.flush()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Project-level runner for the SCAU thesis Word-template workflow."
    )
    parser.add_argument("--project-root", required=True, help="Thesis workspace root.")
    parser.add_argument("--official-template-docx", help="Official template .docx used for frontmatter bootstrap and donor fallback.")
    parser.add_argument("--metadata-file", help="Path to metadata JSON. Supports thesis_metadata.json and generic metadata.json names.")
    parser.add_argument("--docx", help="Base working .docx file.")
    parser.add_argument("--chapter-file", help="Markdown chapter draft to insert.")
    parser.add_argument("--figures-root", help="Root containing figure folders such as 实验结果3-*.")
    parser.add_argument("--figure-number-prefix", help="Only include figure folders matching this prefix, such as 3-.")
    parser.add_argument("--tables-manifest", help="Table manifest JSON path for table-block insertion.")
    parser.add_argument("--references-file", help="Reference source markdown/text file.")
    parser.add_argument("--manifest-output", help="Output path for the generated figure manifest.")
    parser.add_argument("--output", help="Final output .docx path.")
    parser.add_argument(
        "--figure-backend",
        choices=["python-docx", "word-com"],
        default="python-docx",
        help="Backend for figure insertion. Use word-com for large local Word documents on Windows.",
    )
    parser.add_argument(
        "--word-visible",
        action="store_true",
        help="Show the Word window when a Word COM backend is used.",
    )
    parser.add_argument(
        "--chapter-heading",
        help="Normalized chapter heading in Word, such as '3  结果与分析'. Used when trimming template sample body.",
    )
    parser.add_argument("--skip-frontmatter", action="store_true", help="Do not refresh cover, declarations, and abstract frontmatter.")
    parser.add_argument("--skip-generate-figures", action="store_true", help="Do not rerun figure generation scripts.")
    parser.add_argument("--skip-chapter", action="store_true", help="Do not insert chapter Markdown.")
    parser.add_argument("--skip-figures", action="store_true", help="Do not insert figure blocks.")
    parser.add_argument("--skip-tables", action="store_true", help="Do not insert table blocks.")
    parser.add_argument("--skip-references", action="store_true", help="Do not insert references.")
    parser.add_argument(
        "--keep-template-body",
        action="store_true",
        help="Keep the template's sample body chapters instead of trimming them after insertion.",
    )
    return parser.parse_args()


def ensure_sibling_script(name: str) -> Path:
    path = SCRIPT_DIR / name
    if not path.exists():
        raise FileNotFoundError(f"未找到 skill 脚本: {path}")
    return path


def run_step(step_name: str, command: list[str], cwd: Path) -> str:
    result = subprocess.run(
        command,
        check=False,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if result.stdout.strip():
        emit_text(result.stdout.strip())
    if result.returncode != 0:
        recovery_hints = {
            "frontmatter": "检查 thesis_metadata.json、官方模板 docx 路径，以及封面段落锚点是否仍与学校模板一致。",
            "figure-manifest": "检查图目录命名、覆盖配置文件和章节 Markdown 是否存在；必要时先关闭 --run-generators。",
            "chapter": "检查章节 Markdown 的一级标题是否规范，以及工作稿中 donor 样式是否还存在。",
            "tables": "检查表格 manifest、Markdown 表格文件和表题格式是否符合 `# 表x-x ...`。",
            "figures": "检查锚点 regex、图片文件路径，以及是否应传入 --official-template-docx 作为 donor fallback。",
            "references": "检查参考文献源文件格式，确认 `## 中文文献` / `## 英文文献` 结构完整。",
            "contents-finalize": "检查本机是否为 Windows + Word + pywin32 环境，以及目录域是否能正常更新。",
        }
        raise SkillStepError(
            json.dumps(
                {
                    "step": step_name,
                    "returncode": result.returncode,
                    "command": command,
                    "stdout": result.stdout.strip(),
                    "stderr": result.stderr.strip(),
                    "recovery_hint": recovery_hints.get(step_name, "查看 stderr 并从当前步骤重新运行。"),
                },
                ensure_ascii=False,
                indent=2,
            )
        )
    return result.stdout.strip()


def normalized_heading_from_markdown(path: Path) -> str:
    heading_re = re.compile(r"^#\s+(.+?)\s*$", re.MULTILINE)
    text = path.read_text(encoding="utf-8")
    match = heading_re.search(text)
    if not match:
        raise RuntimeError(f"未能从 Markdown 章节草稿识别一级标题: {path}")
    raw = re.sub(r"\s+", " ", match.group(1)).strip()
    chapter_match = re.match(r"^第\s*(\d+)\s*章\s*(.+)$", raw)
    if chapter_match:
        return f"{chapter_match.group(1)}  {chapter_match.group(2).strip()}"
    return raw


def chapter_tag_from_heading(heading: str) -> str:
    match = re.match(r"^(\d+)\s{2,}(.+)$", heading)
    if match:
        return f"第{match.group(1)}章"
    return "章节"


def discover_work_output_dir(project_root: Path) -> Path:
    for candidate in (
        project_root / "论文终稿",
        project_root / "work",
        project_root / "output",
        project_root / "outputs",
    ):
        if candidate.exists() and candidate.is_dir():
            return candidate
    return project_root / "_scau_thesis_output"


def discover_manifest_dir(project_root: Path) -> Path:
    work_dir = discover_work_output_dir(project_root)
    if work_dir.name == "论文终稿":
        return work_dir / "装版清单"
    return work_dir / "manifests"


def discover_template_docx(project_root: Path) -> Path:
    candidates = [
        project_root / "论文撰写规范" / "附件6_格式模板_转存.docx",
        project_root / "论文撰写规范" / "附件6.华南农业大学本科毕业论文（设计）格式模板.docx",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate

    workspace_named = sorted(project_root.rglob("*格式模板*.docx"))
    if workspace_named:
        return workspace_named[0]

    english_named = sorted(project_root.rglob("*template*.docx"))
    if english_named:
        return english_named[0]

    if BUNDLED_TEMPLATE_DOCX.exists():
        return BUNDLED_TEMPLATE_DOCX
    raise FileNotFoundError("未找到可用的华农论文模板 docx。公开仓库首次使用前请先运行 scripts/import_official_2024_assets.py，或显式传入 --official-template-docx。")


def discover_metadata_file(project_root: Path) -> Path:
    direct = project_root / "thesis_metadata.json"
    if direct.exists():
        return direct
    generic = sorted(project_root.rglob("*metadata*.json"))
    if generic:
        return generic[0]
    raise FileNotFoundError("未找到 metadata JSON，请显式传入 --metadata-file。")


def discover_docx_path(project_root: Path) -> Path:
    work_dir = discover_work_output_dir(project_root)
    preferred = [
        work_dir / "毕业论文终稿_工作版.docx",
        work_dir / "scau_thesis_working.docx",
    ]
    for candidate in preferred:
        if candidate.exists():
            return candidate
    return preferred[0]


def discover_figures_root(project_root: Path) -> Path:
    candidates = [
        project_root / "论文草稿",
        project_root / "figures",
        project_root / "images",
        project_root / "assets" / "figures",
        project_root / "图件",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return project_root


def discover_references_file(project_root: Path) -> Path:
    candidates = [
        project_root / "论文草稿" / "最终参考文献著录初稿.md",
        project_root / "references.md",
        project_root / "references.txt",
        project_root / "bibliography.md",
        project_root / "bibliography.txt",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    markdown_matches = sorted(project_root.rglob("*参考文献*.md")) + sorted(project_root.rglob("*reference*.md"))
    if markdown_matches:
        return markdown_matches[0]
    return candidates[0]


def default_manifest_path(project_root: Path, figure_prefix: str | None, heading: str | None) -> Path:
    chapter_number = None
    if figure_prefix:
        chapter_number = figure_prefix.split("-")[0]
    elif heading:
        match = re.match(r"^(\d+)\s", heading)
        if match:
            chapter_number = match.group(1)
    filename = f"第{chapter_number}章图块清单.json" if chapter_number else "图块清单.json"
    return discover_manifest_dir(project_root) / filename


def default_output_path(project_root: Path, heading: str | None) -> Path:
    tag = chapter_tag_from_heading(heading) if heading else "装版"
    work_dir = discover_work_output_dir(project_root)
    if work_dir.name == "论文终稿":
        return work_dir / f"毕业论文终稿_工作版_{tag}装版.docx"
    return work_dir / f"scau_thesis_{tag}_assembled.docx"


def strip_template_body_placeholders(docx_path: Path, chapter_heading: str, output_path: Path) -> Path:
    document = Document(docx_path)
    body_start_re = re.compile(r"^1\s+绪论$")
    target_heading_re = re.compile(rf"^{re.escape(chapter_heading)}$")

    start_index = None
    end_index = None
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if start_index is None and body_start_re.match(text):
            start_index = index
            continue
        if start_index is not None and target_heading_re.match(text):
            end_index = index
            break

    if start_index is None or end_index is None or start_index >= end_index:
        raise RuntimeError("未能识别模板示例正文与新插入章节之间的范围，无法裁掉模板正文。")

    for paragraph in list(document.paragraphs[start_index:end_index]):
        element = paragraph._element
        element.getparent().remove(element)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    args = parse_args()

    project_root = Path(args.project_root).resolve()
    chapter_path = Path(args.chapter_file).resolve() if args.chapter_file else None
    chapter_heading = args.chapter_heading
    if chapter_heading is None and chapter_path is not None:
        chapter_heading = normalized_heading_from_markdown(chapter_path)

    official_template_docx = (
        Path(args.official_template_docx).resolve()
        if args.official_template_docx
        else discover_template_docx(project_root)
    )
    metadata_file = (
        Path(args.metadata_file).resolve()
        if args.metadata_file
        else discover_metadata_file(project_root)
    )

    docx_path = Path(args.docx).resolve() if args.docx else discover_docx_path(project_root)
    if not docx_path.exists() and not args.skip_frontmatter:
        docx_path = official_template_docx

    figures_root = Path(args.figures_root).resolve() if args.figures_root else discover_figures_root(project_root)
    refs_path = (
        Path(args.references_file).resolve()
        if args.references_file
        else discover_references_file(project_root)
    )
    tables_manifest = Path(args.tables_manifest).resolve() if args.tables_manifest else None
    manifest_output = (
        Path(args.manifest_output).resolve()
        if args.manifest_output
        else default_manifest_path(project_root, args.figure_number_prefix, chapter_heading)
    )
    final_output = Path(args.output).resolve() if args.output else default_output_path(project_root, chapter_heading)

    if not args.skip_chapter and chapter_path is None:
        raise ValueError("未跳过章节插入时，必须提供 --chapter-file。")
    if not args.keep_template_body and not args.skip_chapter and not chapter_heading:
        raise ValueError("需要裁掉模板示例正文时，必须能确定章节标题。")

    manifest_output.parent.mkdir(parents=True, exist_ok=True)
    final_output.parent.mkdir(parents=True, exist_ok=True)

    fill_frontmatter = ensure_sibling_script("fill_scau_frontmatter.py")
    generate_manifest = ensure_sibling_script("generate_figure_manifest_from_dirs.py")
    insert_markdown = ensure_sibling_script("insert_markdown_chapter.py")
    insert_tables = ensure_sibling_script("insert_table_blocks.py")
    insert_figures = ensure_sibling_script(
        "insert_figure_blocks_com.py" if args.figure_backend == "word-com" else "insert_figure_blocks.py"
    )
    insert_refs = ensure_sibling_script("insert_reference_batch.py")
    batch_ops = ensure_sibling_script("batch_word_ops.py")

    current_docx = docx_path
    temp_dir = Path(tempfile.mkdtemp(prefix="thesis-word-pipeline-"))

    if not args.skip_frontmatter:
        frontmatter_docx = temp_dir / "frontmatter_filled.docx"
        frontmatter_source = current_docx if current_docx.exists() else official_template_docx
        run_step(
            "frontmatter",
            [
                sys.executable,
                str(fill_frontmatter),
                "--workspace",
                str(project_root),
                "--meta",
                str(metadata_file),
                "--template",
                str(frontmatter_source),
                "--output",
                str(frontmatter_docx),
            ],
            project_root,
        )
        current_docx = frontmatter_docx

    if not args.skip_figures:
        manifest_cmd = [
            sys.executable,
            str(generate_manifest),
            "--figures-root",
            str(figures_root),
            "--output",
            str(manifest_output),
        ]
        if chapter_path is not None:
            manifest_cmd.extend(["--chapter-file", str(chapter_path)])
        if args.figure_number_prefix:
            manifest_cmd.extend(["--number-prefix", args.figure_number_prefix])
        if not args.skip_generate_figures:
            manifest_cmd.append("--run-generators")
        run_step("figure-manifest", manifest_cmd, project_root)

    if not args.skip_chapter:
        chapter_docx = temp_dir / "chapter_inserted.docx"
        run_step(
            "chapter",
            [
                sys.executable,
                str(insert_markdown),
                "--docx",
                str(current_docx),
                "--chapter-file",
                str(chapter_path),
                "--output",
                str(chapter_docx),
            ],
            project_root,
        )
        current_docx = chapter_docx

    if not args.skip_tables and tables_manifest is not None:
        tables_docx = temp_dir / "tables_inserted.docx"
        run_step(
            "tables",
            [
                sys.executable,
                str(insert_tables),
                "--docx",
                str(current_docx),
                "--manifest",
                str(tables_manifest),
                "--fallback-template-docx",
                str(official_template_docx),
                "--output",
                str(tables_docx),
            ],
            project_root,
        )
        current_docx = tables_docx

    if not args.skip_figures:
        figure_docx = temp_dir / "figures_inserted.docx"
        figure_command = [
            sys.executable,
            str(insert_figures),
            "--docx",
            str(current_docx),
            "--manifest",
            str(manifest_output),
            "--fallback-template-docx",
            str(official_template_docx),
            "--output",
            str(figure_docx),
        ]
        if args.figure_backend == "word-com" and args.word_visible:
            figure_command.append("--visible")
        run_step("figures", figure_command, project_root)
        current_docx = figure_docx

    if not args.skip_references:
        refs_docx = temp_dir / "references_inserted.docx"
        run_step(
            "references",
            [
                sys.executable,
                str(insert_refs),
                "--docx",
                str(current_docx),
                "--references-file",
                str(refs_path),
                "--output",
                str(refs_docx),
            ],
            project_root,
        )
        current_docx = refs_docx

    if not args.keep_template_body and not args.skip_chapter:
        stripped_docx = temp_dir / "template_body_cleared.docx"
        current_docx = strip_template_body_placeholders(current_docx, chapter_heading, stripped_docx)

    final_output.write_bytes(current_docx.read_bytes())
    contents_finalized = False
    if sys.platform == "win32":
        finalize_plan = temp_dir / "finalize_contents_plan.json"
        finalized_docx = temp_dir / "contents_finalized.docx"
        finalize_plan.write_text(
            json.dumps(
                [
                    {
                        "action": "finalize_contents",
                        "mode": "full",
                        "update_fields": True,
                    }
                ],
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        run_step(
            "contents-finalize",
            [
                sys.executable,
                str(batch_ops),
                str(final_output),
                str(finalize_plan),
                "--output",
                str(finalized_docx),
            ],
            project_root,
        )
        final_output.write_bytes(finalized_docx.read_bytes())
        contents_finalized = True

    emit_text(
        json.dumps(
            {
                "output": str(final_output),
                "manifest": str(manifest_output) if not args.skip_figures else None,
                "base_docx": str(docx_path),
                "frontmatter_inserted": not args.skip_frontmatter,
                "chapter_inserted": not args.skip_chapter,
                "tables_inserted": tables_manifest is not None and not args.skip_tables,
                "figures_inserted": not args.skip_figures,
                "figure_backend": None if args.skip_figures else args.figure_backend,
                "references_inserted": not args.skip_references,
                "figure_generators_rerun": not args.skip_generate_figures and not args.skip_figures,
                "chapter_heading": chapter_heading,
                "contents_finalized": contents_finalized,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        if isinstance(exc, SkillStepError):
            emit_text(str(exc), stderr=True)
        else:
            emit_text(
                json.dumps(
                    {
                        "step": "runner",
                        "error": str(exc),
                        "recovery_hint": "检查当前参数组合是否属于 frontmatter-only / chapter-only / figure-only / table-only / reference-only 之一。",
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                stderr=True,
            )
        raise SystemExit(1)
