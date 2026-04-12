#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from word_template_utils import (
    apply_three_line_table_format,
    collapse_ws,
    default_output_path,
    donor_key_for_paragraph_text,
    find_heading_donors,
    insert_paragraph_after,
    insert_paragraph_before,
    insert_table_after,
    insert_table_before,
    set_cell_text,
    format_continued_table_caption,
)


TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$")


@dataclass
class TablePayload:
    caption: str
    rows: list[list[str]]
    note: str | None = None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Insert table blocks into a thesis .docx file, including optional automatic continued-table splitting."
    )
    parser.add_argument("--docx", required=True, help="Input .docx path")
    parser.add_argument("--manifest", required=True, help="Table manifest JSON path")
    parser.add_argument("--fallback-template-docx", help="Optional official template .docx used when the working copy lacks table donor paragraphs.")
    parser.add_argument("--output", help="Output .docx path")
    return parser.parse_args()


def parse_md_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and TABLE_SEPARATOR_RE.match(lines[index + 1]) is not None


def parse_table_markdown(path: Path) -> TablePayload:
    lines = path.read_text(encoding="utf-8").splitlines()
    caption = None
    rows: list[list[str]] | None = None
    note_lines: list[str] = []

    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if caption is None and stripped.startswith("#"):
            caption = collapse_ws(stripped.lstrip("#").strip())
            index += 1
            continue
        if rows is None and is_table_start(lines, index):
            header = parse_md_row(lines[index])
            index += 2
            rows = [header]
            while index < len(lines) and "|" in lines[index]:
                rows.append(parse_md_row(lines[index]))
                index += 1
            continue
        if rows is not None and (stripped.startswith("注：") or stripped.startswith("资料来源：")):
            note_lines.append(collapse_ws(stripped))
        index += 1

    if caption is None:
        raise RuntimeError(f"未能从表格 Markdown 中识别表题: {path}")
    if rows is None:
        raise RuntimeError(f"未能从表格 Markdown 中识别 Markdown 表格: {path}")
    note = " ".join(note_lines) if note_lines else None
    return TablePayload(caption=caption, rows=rows, note=note)


def load_manifest(path: Path) -> list[dict]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        raise ValueError("Table manifest must be a JSON array.")
    return data


def find_anchor(document: Document, pattern: str, occurrence: int = 1):
    regex = re.compile(pattern)
    hits = []
    for paragraph in document.paragraphs:
        if regex.search(paragraph.text):
            hits.append(paragraph)
    if len(hits) < occurrence:
        raise RuntimeError(f"Anchor regex {pattern!r} matched {len(hits)} paragraphs, need {occurrence}.")
    return hits[occurrence - 1]


def resolve_payload(entry: dict, base_dir: Path) -> TablePayload:
    if "table_file" in entry:
        table_path = Path(entry["table_file"])
        if not table_path.is_absolute():
            table_path = (base_dir / table_path).resolve()
        payload = parse_table_markdown(table_path)
    else:
        caption = collapse_ws(entry.get("caption", ""))
        rows = entry.get("rows")
        if not caption or not rows:
            raise ValueError("Each table entry must provide either table_file or both caption and rows.")
        payload = TablePayload(caption=caption, rows=rows, note=entry.get("note"))

    if entry.get("caption"):
        payload.caption = collapse_ws(entry["caption"])
    if entry.get("note") is not None:
        payload.note = collapse_ws(entry["note"]) if entry["note"] else None
    return payload


def estimate_row_weight(row_values: list[str]) -> int:
    longest = max((len(collapse_ws(value)) for value in row_values), default=0)
    return max(1, math.ceil(longest / 22))


def split_rows(
    rows: list[list[str]],
    *,
    max_body_rows_per_segment: int | None = None,
    max_segment_weight: int | None = None,
) -> list[list[list[str]]]:
    if len(rows) <= 1:
        return [rows]

    header = rows[0]
    body_rows = rows[1:]
    if max_body_rows_per_segment is None and max_segment_weight is None:
        return [rows]

    segments: list[list[list[str]]] = []
    current_body: list[list[str]] = []
    current_weight = 0

    for row in body_rows:
        row_weight = estimate_row_weight(row)
        if current_body:
            over_row_limit = (
                max_body_rows_per_segment is not None and len(current_body) >= max_body_rows_per_segment
            )
            over_weight_limit = (
                max_segment_weight is not None and current_weight + row_weight > max_segment_weight
            )
            if over_row_limit or over_weight_limit:
                segments.append([header, *current_body])
                current_body = []
                current_weight = 0
        current_body.append(row)
        current_weight += row_weight

    if current_body:
        segments.append([header, *current_body])

    return segments or [rows]


def render_table(table, rows: list[list[str]]) -> None:
    for row_index, row_values in enumerate(rows):
        if row_index >= len(table.rows):
            table.add_row()
        row = table.rows[row_index]
        while len(row.cells) < len(row_values):
            row._tr.add_tc()
        for col_index, value in enumerate(row_values):
            set_cell_text(row.cells[col_index], value)
    apply_three_line_table_format(table)


def build_items(entry: dict, payload: TablePayload) -> list[tuple[str, object]]:
    segments = split_rows(
        payload.rows,
        max_body_rows_per_segment=entry.get("max_body_rows_per_segment"),
        max_segment_weight=entry.get("max_segment_weight"),
    )
    continued_suffix = entry.get("continued_suffix", "（续表）")

    items: list[tuple[str, object]] = [("paragraph", "")]
    for index, rows in enumerate(segments):
        caption = payload.caption if index == 0 else format_continued_table_caption(payload.caption, continued_suffix)
        items.append(("table_caption", caption))
        items.append(("table", rows))
    if payload.note:
        items.append(("note", payload.note))
    items.append(("paragraph", ""))
    return items


def insert_items_after(document: Document, anchor, items: list[tuple[str, object]], donors: dict[str, object]):
    current = anchor
    for kind, payload in items:
        if kind == "table":
            rows = payload
            table = insert_table_after(document, current, rows=len(rows), cols=len(rows[0]))
            render_table(table, rows)
            current = table
            continue
        text = payload
        donor_key = kind if kind in {"table_caption", "note"} else donor_key_for_paragraph_text(str(text))
        current = insert_paragraph_after(current, donors[donor_key], str(text))
    return current


def insert_items_before(document: Document, anchor, items: list[tuple[str, object]], donors: dict[str, object]):
    current = anchor
    for kind, payload in reversed(items):
        if kind == "table":
            rows = payload
            table = insert_table_before(document, current, rows=len(rows), cols=len(rows[0]))
            render_table(table, rows)
            current = table
            continue
        text = payload
        donor_key = kind if kind in {"table_caption", "note"} else donor_key_for_paragraph_text(str(text))
        current = insert_paragraph_before(current, donors[donor_key], str(text))
    return current


def main() -> None:
    args = parse_args()
    docx_path = Path(args.docx).resolve()
    manifest_path = Path(args.manifest).resolve()
    output_path = Path(args.output).resolve() if args.output else default_output_path(docx_path, "_表格插入")

    document = Document(docx_path)
    entries = load_manifest(manifest_path)
    fallback_document = None
    if args.fallback_template_docx:
        fallback_document = Document(Path(args.fallback_template_docx).resolve())
    required_donors = ["body", "table_caption"]
    if any((resolve_payload(entry, manifest_path.parent).note or "").strip() for entry in entries):
        required_donors.append("note")
    donors = find_heading_donors(
        document,
        required=required_donors,
        fallback_document=fallback_document,
    )
    base_dir = manifest_path.parent

    for entry in entries:
        anchor = find_anchor(
            document,
            entry["anchor_regex"],
            occurrence=int(entry.get("occurrence", 1)),
        )
        payload = resolve_payload(entry, base_dir)
        items = build_items(entry, payload)
        position = entry.get("position", "after").lower()
        if position == "after":
            insert_items_after(document, anchor, items, donors)
        elif position == "before":
            insert_items_before(document, anchor, items, donors)
        else:
            raise ValueError(f"Unsupported table insertion position: {position}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(
        json.dumps(
            {"output": str(output_path), "manifest": str(manifest_path), "entries": len(entries)},
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
