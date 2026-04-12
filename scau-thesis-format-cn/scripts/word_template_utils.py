#!/usr/bin/env python3
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Iterator

from docx.document import Document as _Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


TOP_LEVEL_HEADING_RE = re.compile(r"^\d+\s+\S")
TERMINAL_HEADING_RE = re.compile(r"^(参\s*考\s*文\s*献|附录[A-Z]?|致\s*谢)$")
TABLE_CAPTION_RE = re.compile(r"^表\d+(?:[-–]\d+)?(?:（续表）)?\s+\S")
FIGURE_CAPTION_RE = re.compile(r"^图\d+(?:[-–]\d+)?\s+\S")
DEFAULT_BODY_START_RE = re.compile(r"^\d+\s+\S")
INLINE_MARKUP_RE = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)")


def collapse_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def iter_inline_markup_segments(text: str) -> list[tuple[str, bool, bool]]:
    segments: list[tuple[str, bool, bool]] = []
    cursor = 0
    for match in INLINE_MARKUP_RE.finditer(text):
        if match.start() > cursor:
            plain = text[cursor : match.start()]
            if plain:
                segments.append((plain, False, False))
        token = match.group(0)
        if token.startswith("***") and token.endswith("***"):
            payload = token[3:-3]
            if payload:
                segments.append((payload, True, True))
        elif token.startswith("**") and token.endswith("**"):
            payload = token[2:-2]
            if payload:
                segments.append((payload, True, False))
        elif token.startswith("*") and token.endswith("*"):
            payload = token[1:-1]
            if payload:
                segments.append((payload, False, True))
        cursor = match.end()
    if cursor < len(text):
        plain = text[cursor:]
        if plain:
            segments.append((plain, False, False))
    return segments or [("", False, False)]


def normalize_heading_text(text: str) -> str:
    collapsed = collapse_ws(text)
    match = re.match(r"^第\s*(\d+)\s*章\s*(.+)$", collapsed)
    if match:
        return f"{match.group(1)}  {match.group(2).strip()}"
    return collapsed


def normalize_keyword_heading(text: str) -> str:
    return collapse_ws(text).replace(" ", "")


def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def block_element(block: Paragraph | Table):
    return block._p if isinstance(block, Paragraph) else block._tbl


def same_block(left: Paragraph | Table | None, right: Paragraph | Table | None) -> bool:
    if left is None or right is None:
        return False
    return block_element(left) == block_element(right)


def block_text(block: Paragraph | Table) -> str:
    if isinstance(block, Paragraph):
        return collapse_ws(block.text)
    return collapse_ws("\n".join(cell.text for row in block.rows for cell in row.cells))


def remove_block(block: Paragraph | Table) -> None:
    element = block_element(block)
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def copy_run_format(source_run, target_run) -> None:
    if source_run._r.rPr is not None:
        target_r = target_run._r
        existing = target_r.rPr
        if existing is not None:
            target_r.remove(existing)
        target_r.insert(0, deepcopy(source_run._r.rPr))


def clone_paragraph_properties(source: Paragraph, target: Paragraph) -> None:
    if source._p.pPr is None:
        return
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    target._p.insert(0, deepcopy(source._p.pPr))


def copy_paragraph_style(source: Paragraph, target: Paragraph) -> None:
    if source.style is None:
        return
    candidates = []
    name = getattr(source.style, "name", None)
    style_id = getattr(source.style, "style_id", None)
    if name:
        candidates.append(name)
    if style_id:
        candidates.append(style_id)
    for candidate in candidates:
        try:
            target.style = candidate
            return
        except Exception:
            continue


def get_donor_run(paragraph: Paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def donor_font_name(source_run, attr_name: str, fallback: str) -> str:
    if source_run is None:
        return fallback
    value = getattr(source_run.font, attr_name, None)
    if value:
        return str(value)
    name = getattr(source_run.font, "name", None)
    if name:
        return str(name)
    return fallback


def donor_font_size_pt(source_run, fallback_pt: float | None = None) -> float | None:
    if source_run is None:
        return fallback_pt
    size = source_run.font.size
    if size is None:
        return fallback_pt
    try:
        return round(size.pt, 1)
    except Exception:
        return fallback_pt


def apply_run_template_fonts(target_run, donor_run=None, *, east_asia_fallback: str = "宋体", ascii_fallback: str = "Times New Roman") -> None:
    east_asia_font = donor_font_name(donor_run, "name_far_east", east_asia_fallback)
    ascii_font = donor_font_name(donor_run, "name_ascii", ascii_fallback)
    size_pt = donor_font_size_pt(donor_run)
    set_east_asia_font(target_run, east_asia_font=east_asia_font, ascii_font=ascii_font, size_pt=size_pt)


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def populate_paragraph_runs(paragraph: Paragraph, donor: Paragraph, text: str) -> Paragraph:
    clear_paragraph_runs(paragraph)
    donor_run = get_donor_run(donor)
    segments = iter_inline_markup_segments(text)
    for segment_text, force_bold, force_italic in segments:
        run = paragraph.add_run(segment_text)
        if donor_run is not None:
            copy_run_format(donor_run, run)
        apply_run_template_fonts(run, donor_run)
        if force_bold:
            run.bold = True
        if force_italic:
            run.italic = True
    return paragraph


def seed_paragraph_from_donor(target: Paragraph, donor: Paragraph, text: str) -> Paragraph:
    clone_paragraph_properties(donor, target)
    copy_paragraph_style(donor, target)
    return populate_paragraph_runs(target, donor, text)


def insert_paragraph_after(block: Paragraph | Table, donor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    return seed_paragraph_from_donor(paragraph, donor, text)


def insert_paragraph_before(block: Paragraph | Table, donor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addprevious(new_p)
    paragraph = Paragraph(new_p, block._parent)
    return seed_paragraph_from_donor(paragraph, donor, text)


def insert_table_after(document, block: Paragraph | Table, rows: int, cols: int) -> Table:
    table = document.add_table(rows=rows, cols=cols)
    block_element(block).addnext(table._tbl)
    return table


def insert_table_before(document, block: Paragraph | Table, rows: int, cols: int) -> Table:
    table = document.add_table(rows=rows, cols=cols)
    block_element(block).addprevious(table._tbl)
    return table


def find_paragraph_by_regex(document, pattern: str, *, flags: int = 0) -> Paragraph | None:
    regex = re.compile(pattern, flags)
    for paragraph in document.paragraphs:
        if regex.search(collapse_ws(paragraph.text)):
            return paragraph
    return None


def find_paragraph_by_text(document, text: str) -> Paragraph | None:
    wanted = normalize_heading_text(text)
    for paragraph in document.paragraphs:
        current = normalize_heading_text(paragraph.text)
        if current == wanted:
            return paragraph
    return None


def find_next_section_anchor(blocks: list[Paragraph | Table], start_index: int) -> Paragraph | None:
    for block in blocks[start_index + 1 :]:
        if not isinstance(block, Paragraph):
            continue
        text = normalize_heading_text(block.text)
        if TOP_LEVEL_HEADING_RE.match(text) or TERMINAL_HEADING_RE.match(normalize_keyword_heading(text)):
            return block
    return None


def delete_range(document, start_block: Paragraph | Table, end_block: Paragraph | Table | None) -> None:
    removing = False
    for block in list(iter_block_items(document)):
        if same_block(block, start_block):
            removing = True
        if removing and same_block(block, end_block):
            break
        if removing:
            remove_block(block)


def find_body_start_index(document, body_start_regex: str | re.Pattern[str] | None = None) -> int:
    pattern = body_start_regex
    if pattern is None:
        regex = DEFAULT_BODY_START_RE
    elif isinstance(pattern, re.Pattern):
        regex = pattern
    else:
        regex = re.compile(pattern)

    for index, paragraph in enumerate(document.paragraphs):
        text = normalize_heading_text(paragraph.text)
        if "\t" in paragraph.text:
            continue
        if regex.search(text):
            return index
    return 0


def iter_body_paragraphs(document, body_start_regex: str | re.Pattern[str] | None = None) -> Iterator[Paragraph]:
    start_index = find_body_start_index(document, body_start_regex=body_start_regex)
    for paragraph in document.paragraphs[start_index:]:
        yield paragraph


def classify_paragraph_as_donor_key(text: str) -> str | None:
    normalized = collapse_ws(text)
    if not normalized:
        return None
    if "摘" in normalized and "要" in normalized:
        return "abstract_title"
    if re.match(r"^\d+\.\d+\.\d+\.\d+\s+\S", normalized):
        return "heading4"
    if re.match(r"^\d+\.\d+\.\d+\s+\S", normalized):
        return "heading3"
    if re.match(r"^\d+\.\d+\s+\S", normalized):
        return "heading2"
    if re.match(r"^\d+\s+\S", normalize_heading_text(normalized)):
        return "heading1"
    if FIGURE_CAPTION_RE.match(normalized):
        return "figure_caption"
    if TABLE_CAPTION_RE.match(normalized):
        return "table_caption"
    if normalized.startswith("注：") or normalized.startswith("资料来源："):
        return "note"
    if "正文文字" in normalized:
        return "body"
    return None


def find_heading_donors(
    document,
    *,
    required: list[str] | None = None,
    fallback_document=None,
    body_start_regex: str | re.Pattern[str] | None = None,
    allow_body_fallback: bool = True,
) -> dict[str, Paragraph]:
    donors: dict[str, Paragraph] = {}
    search_spaces = [(document, iter_body_paragraphs(document, body_start_regex=body_start_regex))]
    if fallback_document is not None:
        search_spaces.append((fallback_document, iter_body_paragraphs(fallback_document, body_start_regex=body_start_regex)))

    for _, paragraphs in search_spaces:
        for paragraph in paragraphs:
            key = classify_paragraph_as_donor_key(paragraph.text)
            if key and key not in donors:
                donors[key] = paragraph

    if "body" not in donors and allow_body_fallback:
        for paragraph in document.paragraphs:
            if collapse_ws(paragraph.text):
                donors["body"] = paragraph
                break
        if "body" not in donors and fallback_document is not None:
            for paragraph in fallback_document.paragraphs:
                if collapse_ws(paragraph.text):
                    donors["body"] = paragraph
                    break

    required_keys = required or ["heading1", "heading2", "heading3", "heading4", "figure_caption", "table_caption", "note", "body"]
    missing = [key for key in required_keys if key not in donors]
    if missing:
        raise RuntimeError(f"Could not find donor paragraphs for: {', '.join(missing)}")
    return donors


def donor_key_for_paragraph_text(text: str) -> str:
    normalized = collapse_ws(text)
    if normalized.startswith("注：") or normalized.startswith("资料来源："):
        return "note"
    if TABLE_CAPTION_RE.match(normalized):
        return "table_caption"
    if FIGURE_CAPTION_RE.match(normalized):
        return "figure_caption"
    return "body"


def set_east_asia_font(run, east_asia_font: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float | None = None) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt) if size_pt else run.font.size
    run_pr = run._element.get_or_add_rPr()
    run_fonts = run_pr.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_pr.insert(0, run_fonts)
    run_fonts.set(qn("w:ascii"), ascii_font)
    run_fonts.set(qn("w:hAnsi"), ascii_font)
    run_fonts.set(qn("w:cs"), ascii_font)
    run_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt:
        half_points = str(int(size_pt * 2))
        sz = run_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            run_pr.append(sz)
        sz.set(qn("w:val"), half_points)
        sz_cs = run_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            run_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), half_points)


def set_hanging_indent_chars(paragraph: Paragraph, chars: float = 2.0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    char_units = str(int(chars * 100))
    twips = str(int(chars * 240))
    ind.set(qn("w:leftChars"), char_units)
    ind.set(qn("w:hangingChars"), char_units)
    ind.set(qn("w:left"), twips)
    ind.set(qn("w:hanging"), twips)


def replace_paragraph_text_from_donor(paragraph: Paragraph, donor: Paragraph, text: str) -> Paragraph:
    donor_ppr = deepcopy(donor._p.pPr) if donor._p.pPr is not None else None
    donor_style = donor.style
    clear_paragraph_runs(paragraph)
    if paragraph._p.pPr is not None:
        paragraph._p.remove(paragraph._p.pPr)
    if donor_ppr is not None:
        paragraph._p.insert(0, donor_ppr)
    if donor_style is not None:
        paragraph.style = donor_style
    return populate_paragraph_runs(paragraph, donor, text)


def set_cell_text(cell, text: str, *, east_asia_font: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 10.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.first_line_indent = None
    clear_paragraph_runs(paragraph)
    for segment_text, force_bold, force_italic in iter_inline_markup_segments(text):
        run = paragraph.add_run(segment_text)
        set_east_asia_font(run, east_asia_font=east_asia_font, ascii_font=ascii_font, size_pt=size_pt)
        if force_bold:
            run.bold = True
        if force_italic:
            run.italic = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_continued_table_caption(caption: str, continued_suffix: str = "（续表）") -> str:
    normalized = collapse_ws(caption)
    match = re.match(r"^(表\d+(?:[-–]\d+)?)(?:（续表）)?(\s+.+)$", normalized)
    if not match:
        return f"{normalized}{continued_suffix}"
    return f"{match.group(1)}{continued_suffix}{match.group(2)}"


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge, edge_data in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_three_line_table_format(table: Table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )

    if table.rows:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 12, "space": 0, "color": "000000"},
                bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
            )
        for cell in table.rows[-1].cells:
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": 12, "space": 0, "color": "000000"},
            )


def default_output_path(path: Path, suffix: str) -> Path:
    return path.with_name(f"{path.stem}{suffix}{path.suffix}")
